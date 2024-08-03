import os
import streamlit as st
from groq import Groq
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import PyPDF2

# Groq client setup
client = Groq(api_key="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")

# Function to analyze and suggest improvements
def analyze_and_suggest(jd, resume):
    jd_analysis = client.chat.completions.create(
        messages=[{"role": "user", "content": f"Extract skills and provide an analysis for the following job description: {jd}"}],
        model="llama3-8b-8192",
    )
    jd_skills = extract_skills_from_response(jd_analysis)

    resume_analysis = client.chat.completions.create(
        messages=[{"role": "user", "content": f"Extract skills and provide an analysis for the following resume: {resume}"}],
        model="llama3-8b-8192",
    )
    resume_skills = extract_skills_from_response(resume_analysis)

    suggestions_analysis = client.chat.completions.create(
        messages=[{"role": "user", "content": f"Provide suggestions for aligning the following resume with this job description: {resume} {jd}"}],
        model="llama3-8b-8192",
    )
    suggestions = suggestions_analysis.choices[0].message.content

    missing_skills = set(jd_skills) - set(resume_skills)
    
    return {
        "skills_to_add": list(missing_skills),
        "alignment_tips": suggestions
    }

# Function to extract skills from response
def extract_skills_from_response(response):
    content = response.choices[0].message.content
    skills = content.split(", ")  # Placeholder for actual skill extraction logic
    return skills

# Function to read file content with different encodings
def read_file(file, file_type):
    if file_type == "txt":
        try:
            return file.read().decode("utf-8")
        except UnicodeDecodeError:
            return file.read().decode("ISO-8859-1")
    elif file_type == "pdf":
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    elif file_type == "docx":
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])

# Function to create the resume document
def create_resume(text, file_path):
    doc = Document()
    set_margins(doc, 0.189, 0.189, 0.276, 0.276)

    # Adding Contact Information
    contact_heading = doc.add_paragraph()
    contact_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = contact_heading.add_run("Rushitha Anugu | Ph: +19408435023 | Email: rushith1728@gmail.com | ")
    run.bold = True
    run.font.size = Pt(12)

    # Add LinkedIn hyperlink
    add_hyperlink(contact_heading, 'https://www.linkedin.com/in/rushitha/', 'LinkedIn')
    run = contact_heading.add_run(" | ")

    # Add GitHub hyperlink
    add_hyperlink(contact_heading, 'https://github.com/rushitha', 'GitHub')
    run = contact_heading.add_run(" | Houston, TX")

    # Add sections and paragraphs
    add_resume_content(doc, text)

    doc.save(file_path)
    print(f"Document saved at: {file_path}")

def add_resume_content(doc, text):
    sections = text.split('\n\n')

    for section in sections:
        lines = section.split('\n')
        if len(lines) > 0:
            heading = lines[0]
            add_heading(doc, heading, level=1)
            for line in lines[1:]:
                add_paragraph(doc, line, bullet=line.startswith("•"))

def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    if level == 1:
        run.font.size = Pt(14)
    return paragraph

def add_paragraph(doc, text, bullet=False):
    paragraph = doc.add_paragraph(style='List Bullet' if bullet else None)
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    return paragraph

def set_single_line_spacing(paragraph):
    """ Set single line spacing for a paragraph """
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "240")  # line value in twentieths of a point (240 Twips = 12 pt)
    pPr.append(spacing)

def set_margins(doc, top, bottom, right, left):
    """ Set custom page margins """
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

def add_hyperlink(paragraph, url, text):
    """Function to add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a run and add it to the hyperlink
    r = OxmlElement('w:r')
    hyperlink.append(r)
    rPr = OxmlElement('w:rPr')  # Create run properties
    r.append(rPr)
    style = OxmlElement('w:rStyle')
    style.set(qn('w:val'), 'Hyperlink')
    rPr.append(style)

    t = OxmlElement('w:t')  # Create text element
    t.text = text
    r.append(t)
    paragraph._p.append(hyperlink)

# Initialize session state for storing suggestions
if "suggestions" not in st.session_state:
    st.session_state["suggestions"] = None

# Sidebar for navigation
st.sidebar.title("Navigation")
page = st.sidebar.radio("Go to", ["Resume Analyzer", "View Suggestions", "Update Resume", "Generate Python Code", "Placeholder 1", "Placeholder 2"])

# Resume Analyzer Interface
if page == "Resume Analyzer":
    st.title("Rushitha's Job Description and Resume Analyzer")
    st.write("Upload a job description and your resume or paste the text to get suggestions on how to align your resume better with the job description.")

    jd_mode = st.radio("Job Description Input Mode", ("Upload Document", "Paste Text"))
    if jd_mode == "Upload Document":
        jd_file = st.file_uploader("Upload Job Description", type=["txt", "pdf", "docx"])
        if jd_file:
            jd_content = read_file(jd_file, jd_file.name.split(".")[-1])
        else:
            jd_content = ""
    else:
        jd_content = st.text_area("Paste Job Description Text Here")

    resume_mode = st.radio("Resume Input Mode", ("Upload Document", "Paste Text"))
    if resume_mode == "Upload Document":
        resume_file = st.file_uploader("Upload Resume", type=["txt", "pdf", "docx"])
        if resume_file:
            resume_content = read_file(resume_file, resume_file.name.split(".")[-1])
        else:
            resume_content = ""
    else:
        resume_content = st.text_area("Paste Resume Text Here")

    if st.button("Submit"):
        if jd_content and resume_content:
            st.session_state["suggestions"] = analyze_and_suggest(jd_content, resume_content)
            st.write("Analysis completed! Please go to the 'View Suggestions' page to see the results.")
        else:
            st.write("Please provide both the job description and the resume.")

# View Suggestions Interface
elif page == "View Suggestions":
    st.title("View Suggestions")
    if st.session_state["suggestions"]:
        suggestions = st.session_state["suggestions"]
        st.write("Skills to Add:")
        st.write(suggestions["skills_to_add"])
        st.write("Alignment Tips:")
        st.write(suggestions["alignment_tips"])
    else:
        st.write("No suggestions available. Please go to the 'Resume Analyzer' page to generate suggestions.")

# Update Resume Interface
elif page == "Update Resume":
    st.title("Update Resume")
    st.write("Update the resume text below and download the updated resume as a .docx file.")

    # Pre-filled resume content
    resume_text = """
Rushitha Anugu
Ph: +19408435023 | email: rushith1728@gmail.com | LinkedIn | GitHub | Houston, TX

WORK EXPERIENCE
Salesforce Developer | Capgemini Technology Services India Ltd | Hyderabad, India                                     May 2019 – Aug 2022
• Managed user setup, roles, profiles, permissions, public groups, OWD, and sharing rules, improving user management efficiency by 30%.
• Performed regular Salesforce system maintenance including security reviews, health checks, and optimizer reviews, ensuring 99.9% system uptime.
• Developed and maintained comprehensive technical documentation for Salesforce configurations and customizations.
• Developed custom applications using Apex, Visualforce, and Lightning Components, reducing manual processes by 25%.
• Created and managed custom objects, fields, workflows, process builder flows, and approval processes, enhancing process automation by 20%.
• Integrated Salesforce with external systems using REST and SOAP APIs, streamlining data flow and reducing data entry time by 40%.
• Implemented batch and schedule apex classes to handle large datasets, optimizing data processing tasks by 35%.
• Deployed code across multiple environments using Change Sets, ensuring smooth and error-free deployments.
• Led a team to implement a new Salesforce-based CRM system, resulting in a 15% increase in sales team productivity.
• Conducted user training sessions, improving user adoption and satisfaction scores by 25%.

Graduate Student Assistant | University of North Texas| Denton, TX                                                               Sep 2023 – May 2024
• Supervised and trained a team of student staff, coordinating schedules and enforcing policies to maintain an orderly and efficient facility.
• Implemented health and safety protocols, reducing incidents and promoting a culture of safety within the facility.

PROJECTS
Salesforce Community Portal Implementation | Jan 2022 – Apr 2022
• Developed a customer community portal using Salesforce Community Cloud, improving customer engagement and support efficiency by 30%.
• Customized community pages using Lightning Components and Visualforce, enhancing user experience and functionality.
• Integrated the community portal with external systems for seamless data flow using REST APIs, reducing manual data entry by 25%.

Location-Based Taxi Aggregator and Selector | Dec 2023 – May 2024
• Developed and maintained a geo-aware database using MongoDB Atlas to store real-time taxi location data within the defined boundaries of downtown Houston.
• Designed and implemented backend services using Python Flask and AWS ECS, facilitating efficient ingestion and processing of taxi location updates and user requests.
• Utilized Google Maps API for real-time map visualization, enabling accurate proximity matching between users and available taxis.
• Deployed the system on AWS infrastructure leveraging EC2 instances, AWS CLI, and Docker for containerization and scalable deployment, ensuring reliable performance during peak usage times.

SKILLS
Salesforce Development: Apex, Visualforce, Lightning Components, Salesforce APIs (REST, SOAP), Process Builder, Flow, Workflow Rules, Custom Objects, Custom Fields, Batch and Schedule Apex Classes, Change Sets
Programming Languages: Python, R, Java, C++, JavaScript, HTML, CSS
Machine Learning / Data Science: Linear and Logistic Regression, Decision Trees, XGBoost, Random Forest, Neural Networks (NLP, LSTM, RNN, CNN)
Data Management: SQL, NoSQL (MongoDB, DynamoDB), PySpark, Power BI, Excel
Tools and Technologies: Git, GitHub, Docker, Kubernetes, AWS (S3, Lambda, Redshift, EC2, EMR), Apache Spark, Hadoop, Jupyter Notebook, TensorFlow, Keras

EDUCATION
University of North Texas | Denton, TX                                                                                                                       Aug 2022 – May 2024
Master’s in Computer Science and Engineering                                                                                      
Jawaharlal Nehru Technological University | Nizamabad, Hyderabad                                                                June 2015 – May 2019
Bachelor’s in Computer Science and Engineering
    """
    
    updated_resume_text = st.text_area("Edit Resume Text Here", resume_text.strip(), height=600)
    
    if st.button("Generate Updated Resume"):
        save_path = os.path.join(os.getcwd(), "Updated_Resume.docx")
        create_resume(updated_resume_text, save_path)
        with open(save_path, "rb") as file:
            st.download_button(
                label="Download Updated Resume",
                data=file,
                file_name="Updated_Resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        st.session_state["updated_resume_text"] = updated_resume_text

# Generate Python Code Interface
elif page == "Generate Python Code":
    st.title("Generate Python Code")
    st.write("This code will generate the updated resume as a .docx file.")

    if "updated_resume_text" in st.session_state:
        updated_resume_text = st.session_state["updated_resume_text"]
        code = f"""
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def set_single_line_spacing(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "240")
    pPr.append(spacing)

def set_margins(doc, top, bottom, right, left):
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    r = OxmlElement('w:r')
    hyperlink.append(r)
    rPr = OxmlElement('w:rPr')
    r.append(rPr)
    style = OxmlElement('w:rStyle')
    style.set(qn('w:val'), 'Hyperlink')
    rPr.append(style)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    paragraph._p.append(hyperlink)

doc = Document()
set_margins(doc, 0.189, 0.189, 0.276, 0.276)

contact_heading = doc.add_paragraph()
contact_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = contact_heading.add_run("Rushitha Anugu | Ph: +19408435023 | Email: rushith1728@gmail.com | ")
add_hyperlink(contact_heading, 'https://www.linkedin.com/in/rushitha/', 'LinkedIn')
run = contact_heading.add_run(" | ")
add_hyperlink(contact_heading, 'https://github.com/rushitha', 'GitHub')
run = contact_heading.add_run(" | Houston, TX")

sections = updated_resume_text.split('\\n\\n')
for section in sections:
    if section.startswith("WORK EXPERIENCE") or section.startswith("PROJECTS") or section.startswith("SKILLS") or section.startswith("EDUCATION"):
        add_heading(doc, section.split('\\n')[0], level=1)
    else:
        add_paragraph(doc, section)

def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    if level == 1:
        run.font.size = Pt(14)
    return paragraph

def add_paragraph(doc, text, bullet=False):
    paragraph = doc.add_paragraph(style='List Bullet' if bullet else None)
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    return paragraph

doc.save('Updated_Resume.docx')
print('Document saved as Updated_Resume.docx')
        """
        st.code(code, language="python")
    else:
        st.write("Please update your resume in the 'Update Resume' section first.")

# Placeholder 1 Interface
elif page == "Placeholder 1":
    st.title("Execute Custom Code")
    st.write("Paste your Python code here to generate a .docx file.")

    custom_code = st.text_area("Paste Python Code Here", height=400)

    updated_resume_text = st.text_area("Paste Updated Resume Text Here", height=200)
    
    if st.button("Run Code"):
        save_path = os.path.join(os.getcwd(), "Generated_Resume.docx")
        exec_globals = {
            "Document": Document,
            "Inches": Inches,
            "Pt": Pt,
            "qn": qn,
            "OxmlElement": OxmlElement,
            "WD_PARAGRAPH_ALIGNMENT": WD_PARAGRAPH_ALIGNMENT,
            "set_single_line_spacing": set_single_line_spacing,
            "set_margins": set_margins,
            "add_hyperlink": add_hyperlink,
            "add_heading": add_heading,
            "add_paragraph": add_paragraph,
            "save_path": save_path,
            "updated_resume_text": updated_resume_text
        }
        try:
            exec(custom_code, exec_globals)
            if os.path.exists(save_path):
                st.write(f"Document successfully generated at {save_path}")
                with open(save_path, "rb") as file:
                    st.download_button(
                        label="Download Generated Resume",
                        data=file,
                        file_name="Generated_Resume.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.write("The code did not generate a file at the expected path.")
        except Exception as e:
            st.write(f"Error: {e}")

# Placeholder 2 Interface
elif page == "Placeholder 2":
    st.title("Placeholder 2")
    st.write("This is a placeholder for future functionality.")
